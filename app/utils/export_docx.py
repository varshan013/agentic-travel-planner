from docx import Document

def export_to_docx(text: str, filename="travel_plan.docx"):
    doc = Document()
    doc.add_heading("Agentic AI Travel Plan", level=1)

    for line in text.split("\n"):
        doc.add_paragraph(line)

    doc.save(filename)
    return filename
